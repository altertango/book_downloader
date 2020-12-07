import urllib.request
from bs4 import BeautifulSoup
import os
from time import sleep
import docx

novel_url = "https://fastnovel.net/abuse-of-magic-749/chapter-290528.html"

novel_name=novel_url.split('/')[3].replace('-',' ').title()
doc = docx.Document() 
doc.add_heading(novel_name, 0)

def get_chapter(chapter):
    url=chapter
    print(url)
    req = urllib.request.Request(
        url, 
        data=None, 
        headers={
            'dnt': '1',
            'upgrade-insecure-requests': '1',
            'user-agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_4) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/83.0.4103.61 Safari/537.36',
            'accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
            'sec-fetch-site': 'none',
            'sec-fetch-mode': 'navigate',
            'sec-fetch-user': '?1',
            'sec-fetch-dest': 'document',
            'accept-language': 'en-GB,en-US;q=0.9,en;q=0.8',
        }
    )
    f = urllib.request.urlopen(req)
    soup = BeautifulSoup(f.read(), 'html.parser')
    chapter_name = soup.find_all("h1",{'class': 'episode-name'})[0].text
    buttons = soup.find_all("a",{'class': 'btn btn-success'})
    prev_chapoter,next_chapter =[i["href"] for i in buttons]
    chapter_text = soup.find("div", {"id": 'chapter-body'})
    r=[]
    for p in chapter_text.findAll('p'):
        #print(''.join(p.findAll(text=True)))
        r.append(''.join(p.findAll(text=True)))
    return chapter_name, r, prev_chapoter,next_chapter








it=0
chapters=[]
n_c=novel_url
while (it<300 and not n_c in chapters):
    print 
    chapters.append(n_c)
    c_n, text, p_c,n_c = get_chapter(n_c)
    n_c="https://fastnovel.net/"+n_c
    doc.add_heading(c_n, 1)
    [doc.add_paragraph(t) for t in text]
    doc.add_page_break()
    it+=1
    sleep(.5)

    
doc.save(novel_name+'.docx') 
